from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable

import pandas as pd
import streamlit as st
from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_core.messages import BaseMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS
from langchain_huggingface import ChatHuggingFace, HuggingFaceEmbeddings, HuggingFaceEndpoint


load_dotenv()
if os.getenv("HUGGINGFACEHUB_API_TOKEN") and not os.getenv("HF_TOKEN"):
    os.environ["HF_TOKEN"] = os.getenv("HUGGINGFACEHUB_API_TOKEN", "")

UPLOAD_DIR = Path(".streamlit_uploads")
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv"}


def get_embedding_model() -> HuggingFaceEmbeddings:
    model_name = os.getenv("HF_EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")
    return HuggingFaceEmbeddings(model=model_name)


def get_llm() -> ChatHuggingFace:
    repo_id = os.getenv("HF_LLM_REPO_ID", "meta-llama/Meta-Llama-3-8B-Instruct")
    token = os.getenv("HUGGINGFACEHUB_API_TOKEN")
    if not token:
        raise RuntimeError("Set HUGGINGFACEHUB_API_TOKEN in your environment or .env file.")

    endpoint = HuggingFaceEndpoint(
        repo_id=repo_id,
        huggingfacehub_api_token=token,
        task="conversational",
        temperature=0.2,
        max_new_tokens=700,
        top_p=0.9,
    )
    return ChatHuggingFace(llm=endpoint)


def response_text(response: object) -> str:
    if isinstance(response, BaseMessage):
        return str(response.content)
    return str(response)


def save_uploaded_files(uploaded_files: Iterable[st.runtime.uploaded_file_manager.UploadedFile]) -> list[Path]:
    UPLOAD_DIR.mkdir(exist_ok=True)
    saved_paths: list[Path] = []

    for uploaded_file in uploaded_files:
        suffix = Path(uploaded_file.name).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            continue

        destination = UPLOAD_DIR / uploaded_file.name
        destination.write_bytes(uploaded_file.getbuffer())
        saved_paths.append(destination)

    return saved_paths


def load_document(path: Path) -> list[Document]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        docs = PyPDFLoader(str(path)).load()
    elif suffix == ".docx":
        docs = load_word_document(path)
    elif suffix in {".xlsx", ".xls", ".csv"}:
        docs = load_tabular_document(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    for doc in docs:
        doc.metadata["source"] = path.name
    return docs


def load_word_document(path: Path) -> list[Document]:
    document = DocxDocument(path)
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    table_blocks: list[str] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            table_blocks.append(f"Table {table_index}\n" + "\n".join(rows))

    content = "\n\n".join(paragraphs + table_blocks)
    return [Document(page_content=content, metadata={"source": path.name})] if content else []


def load_tabular_document(path: Path) -> list[Document]:
    if path.suffix.lower() == ".csv":
        frames = {"CSV": pd.read_csv(path)}
    else:
        frames = pd.read_excel(path, sheet_name=None)

    documents: list[Document] = []
    for sheet_name, frame in frames.items():
        cleaned = frame.dropna(how="all").fillna("")
        if cleaned.empty:
            continue

        content = cleaned.to_markdown(index=False)
        documents.append(
            Document(
                page_content=content,
                metadata={"source": path.name, "sheet_name": str(sheet_name)},
            )
        )

    return documents


def load_documents(paths: Iterable[Path]) -> list[Document]:
    documents: list[Document] = []
    for path in paths:
        documents.extend(load_document(path))
    return documents


def split_documents(documents: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=900,
        chunk_overlap=150,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    return splitter.split_documents(documents)


def build_vector_store(chunks: list[Document]) -> FAISS:
    return FAISS.from_documents(chunks, get_embedding_model())


def citation_label(doc: Document) -> str:
    source = doc.metadata.get("source", "Unknown source")
    page = doc.metadata.get("page")
    sheet = doc.metadata.get("page_name") or doc.metadata.get("sheet_name")

    details: list[str] = []
    if page is not None:
        details.append(f"page {int(page) + 1}")
    if sheet:
        details.append(f"sheet {sheet}")

    return f"{source} ({', '.join(details)})" if details else source


def render_citations(documents: list[Document]) -> None:
    if not documents:
        return

    st.subheader("Citations")
    seen: set[str] = set()
    for doc in documents:
        label = citation_label(doc)
        if label in seen:
            continue
        seen.add(label)
        snippet = " ".join(doc.page_content.split())[:450]
        st.markdown(f"**{label}**")
        st.caption(snippet)


def answer_question(question: str, vector_store: FAISS) -> tuple[str, list[Document]]:
    llm = get_llm()
    source_docs = vector_store.similarity_search(question, k=5)
    context = "\n\n".join(
        f"Source: {citation_label(doc)}\n{doc.page_content}" for doc in source_docs
    )

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "You are an enterprise knowledge assistant. Answer only from the provided context. "
                "If the answer is not in the context, say you do not know. Include concise source references "
                "inside the answer using the source names when possible.\n\nContext:\n{context}",
            ),
            ("human", "{input}"),
        ]
    )
    answer = (prompt | llm).invoke({"input": question, "context": context})
    return response_text(answer), source_docs


def summarize_collection(vector_store: FAISS) -> tuple[str, list[Document]]:
    docs = vector_store.similarity_search(
        "main topics, key decisions, policy rules, risks, metrics, owners, dates, and action items",
        k=8,
    )
    context = "\n\n".join(
        f"Source: {citation_label(doc)}\n{doc.page_content}" for doc in docs
    )

    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Summarize the enterprise document collection from the provided excerpts. "
                "Use short sections for overview, key points, decisions, risks, and action items. "
                "Mention source names where useful.",
            ),
            ("human", "{context}"),
        ]
    )
    answer = (prompt | get_llm()).invoke({"context": context})
    return response_text(answer), docs


def initialize_state() -> None:
    st.session_state.setdefault("vector_store", None)
    st.session_state.setdefault("document_count", 0)
    st.session_state.setdefault("chunk_count", 0)


def main() -> None:
    st.set_page_config(
        page_title="Enterprise Knowledge Assistant",
        layout="wide",
    )
    initialize_state()

    st.title("Enterprise Knowledge Assistant")
    st.caption("RAG-powered document search, summarization, and cited answers for internal knowledge.")

    with st.sidebar:
        st.header("Knowledge Base")
        uploaded_files = st.file_uploader(
            "Upload PDFs, Word files, Excel sheets, or CSVs",
            type=["pdf", "docx", "xlsx", "xls", "csv"],
            accept_multiple_files=True,
        )

        if st.button("Build Knowledge Base", type="primary", use_container_width=True):
            if not uploaded_files:
                st.warning("Upload at least one supported document first.")
            else:
                with st.status("Indexing documents...", expanded=True) as status:
                    saved_paths = save_uploaded_files(uploaded_files)
                    st.write(f"Saved {len(saved_paths)} files")
                    documents = load_documents(saved_paths)
                    st.write(f"Loaded {len(documents)} document sections")
                    chunks = split_documents(documents)
                    st.write(f"Created {len(chunks)} searchable chunks")
                    st.session_state.vector_store = build_vector_store(chunks)
                    st.session_state.document_count = len(documents)
                    st.session_state.chunk_count = len(chunks)
                    status.update(label="Knowledge base ready", state="complete")

        st.divider()
        st.metric("Documents", st.session_state.document_count)
        st.metric("Chunks", st.session_state.chunk_count)

    vector_store = st.session_state.vector_store

    if vector_store is None:
        st.info("Upload documents and build the knowledge base to begin.")
        return

    question_tab, summary_tab = st.tabs(["Ask", "Summarize"])

    with question_tab:
        question = st.text_input("Question", placeholder="What is the remote work policy?")
        if st.button("Ask Assistant", use_container_width=True):
            if not question.strip():
                st.warning("Enter a question first.")
            else:
                try:
                    with st.spinner("Retrieving context and generating answer..."):
                        answer, source_docs = answer_question(question, vector_store)
                    st.subheader("Answer")
                    st.write(answer)
                    render_citations(source_docs)
                except Exception as exc:
                    st.error(str(exc))

    with summary_tab:
        if st.button("Summarize Documents", use_container_width=True):
            try:
                with st.spinner("Creating summary..."):
                    summary, source_docs = summarize_collection(vector_store)
                st.subheader("Summary")
                st.write(summary)
                render_citations(source_docs)
            except Exception as exc:
                st.error(str(exc))


if __name__ == "__main__":
    main()
